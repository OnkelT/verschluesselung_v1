import os
import random
from fpdf import FPDF
from docx import Document  # Für die Erstellung von .docx-Dokumenten

def create_docx(file_path):
    """
    Erzeugt ein minimales .docx-Dokument mit einem Titel und einem kurzen Absatz.
    """
    doc = Document()
    doc.add_heading("Generiertes Dokument", level=1)
    doc.add_paragraph("Dies ist ein automatisch generiertes Dokument.")
    doc.save(file_path)

def create_pdf(file_path):
    """
    Erzeugt ein minimales PDF, das von PDF-Readern erkannt wird.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Eine leere Zelle sorgt für einen korrekten PDF-Header.
    pdf.cell(0, 10, txt="", ln=1)
    pdf.output(file_path)

def main():
    # Ermittle den Ordner, in dem sich diese Datei befindet.
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # Erstelle einen Zielordner "erzeugte_dateien" im gleichen Verzeichnis wie das Script.
    output_folder = os.path.join(script_dir, "erzeugte_dateien")
    os.makedirs(output_folder, exist_ok=True)

    num_files = 200
    print(f"Erstelle {num_files} .docx-Dokumente und {num_files} PDF-Dateien im Ordner:\n{output_folder}\n")

    for i in range(1, num_files + 1):
        # Generiere eine zufällige 6-stellige Rechnungsnummer.
        rechnungsnummer = str(random.randint(100000, 999999))
        id_str = str(i)
        base_name = f"({rechnungsnummer},{id_str})"

        docx_file = os.path.join(output_folder, base_name + ".docx")
        pdf_file = os.path.join(output_folder, base_name + ".pdf")

        try:
            create_docx(docx_file)
            print(f".docx-Dokument erstellt: {docx_file}")
        except Exception as e:
            print(f"Fehler beim Erstellen von {docx_file}: {e}")

        try:
            create_pdf(pdf_file)
            print(f"PDF erstellt: {pdf_file}")
        except Exception as e:
            print(f"Fehler beim Erstellen von {pdf_file}: {e}")

    print("\n✅ Fertig: 200 .docx-Dokumente und 200 PDFs wurden erstellt.")

if __name__ == "__main__":
    main()

